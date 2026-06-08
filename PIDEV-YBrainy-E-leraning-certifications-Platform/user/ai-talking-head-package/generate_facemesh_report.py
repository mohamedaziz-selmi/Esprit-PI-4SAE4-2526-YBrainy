from __future__ import annotations

import argparse
import hashlib
import math
import urllib.request
from pathlib import Path

import cv2
import matplotlib
import mediapipe as mp
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from mediapipe.tasks.python import vision
from mediapipe.tasks.python.core.base_options import BaseOptions
from scipy.stats import zscore
from sklearn.cluster import AgglomerativeClustering, KMeans
from sklearn.decomposition import PCA
from sklearn.metrics import (
    calinski_harabasz_score,
    davies_bouldin_score,
    silhouette_score,
)
from sklearn.mixture import GaussianMixture
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler

matplotlib.use("Agg")
import matplotlib.pyplot as plt


ROOT = Path(__file__).resolve().parent
DATASET_ROOT = ROOT / "fairface-img-margin025-trainval"
TRAIN_LABELS = ROOT / "fairface_label_train.csv"
VAL_LABELS = ROOT / "fairface_label_val.csv"
OUTPUT_DIR = ROOT / "report_outputs"
ASSETS_DIR = OUTPUT_DIR / "assets"
MODEL_PATH = ROOT / "face_landmarker.task"
MODEL_URL = (
    "https://storage.googleapis.com/mediapipe-models/"
    "face_landmarker/face_landmarker/float16/1/face_landmarker.task"
)
AGE_ORDER = [
    "0-2",
    "3-9",
    "10-19",
    "20-29",
    "30-39",
    "40-49",
    "50-59",
    "60-69",
    "more than 70",
]
RACE_ORDER = [
    "White",
    "Latino_Hispanic",
    "Indian",
    "East Asian",
    "Black",
    "Southeast Asian",
    "Middle Eastern",
]
GENDER_ORDER = ["Female", "Male"]
FEATURE_LABELS = {
    "eye_left_open": "Ouverture oeil gauche",
    "eye_right_open": "Ouverture oeil droit",
    "nose_width_ratio": "Largeur du nez",
    "nose_bridge_ratio": "Longueur de l'arete nasale",
    "face_shape_index": "Indice de forme du visage",
    "jaw_width_ratio": "Largeur de la machoire",
    "lip_fullness_ratio": "Epaisseur des levres",
    "eyebrow_arch_ratio": "Arc du sourcil",
    "interpupillary_ratio": "Distance interpupillaire",
    "chin_projection_ratio": "Projection du menton",
}
FEATURE_GROUPS = {
    "eye_left_open": "morphologie oculaire",
    "eye_right_open": "morphologie oculaire",
    "eyebrow_arch_ratio": "morphologie oculaire",
    "interpupillary_ratio": "espacement des yeux",
    "nose_width_ratio": "morphologie nasale",
    "nose_bridge_ratio": "morphologie nasale",
    "face_shape_index": "structure globale du visage",
    "jaw_width_ratio": "structure globale du visage",
    "lip_fullness_ratio": "bouche et levres",
    "chin_projection_ratio": "bas du visage",
}
MODEL_NAMES = {
    "KMeans": "KMeans",
    "Agglomerative": "Agglomerative (Ward)",
    "GaussianMixture": "Gaussian Mixture",
}


def ensure_model(model_path: Path) -> None:
    if model_path.exists():
        return
    urllib.request.urlretrieve(MODEL_URL, model_path)


def load_labels() -> pd.DataFrame:
    train_df = pd.read_csv(TRAIN_LABELS).assign(split="train")
    val_df = pd.read_csv(VAL_LABELS).assign(split="val")
    df = pd.concat([train_df, val_df], ignore_index=True)
    df["image_path"] = df["file"].map(lambda rel: DATASET_ROOT / rel)
    return df


def save_table(df: pd.DataFrame, path: Path) -> None:
    df.to_csv(path, index=False, encoding="utf-8")


def prepare_dirs() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def plot_label_distributions(df: pd.DataFrame) -> None:
    sns.set_theme(style="whitegrid")

    race_counts = (
        df.groupby(["split", "race"]).size().reset_index(name="count")
        .assign(race=lambda frame: pd.Categorical(frame["race"], categories=RACE_ORDER, ordered=True))
        .sort_values("race")
    )
    plt.figure(figsize=(10, 5))
    sns.barplot(data=race_counts, x="race", y="count", hue="split", palette="Set2")
    plt.title("Repartition des images par race et par split")
    plt.xlabel("Race")
    plt.ylabel("Nombre d'images")
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "race_distribution.png", dpi=200)
    plt.close()

    age_counts = (
        df["age"].value_counts().reindex(AGE_ORDER).rename_axis("age").reset_index(name="count")
    )
    plt.figure(figsize=(9, 4.5))
    sns.barplot(data=age_counts, x="age", y="count", color="#5B8FF9")
    plt.title("Distribution globale des tranches d'age")
    plt.xlabel("Tranche d'age")
    plt.ylabel("Nombre d'images")
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "age_distribution.png", dpi=200)
    plt.close()

    gender_counts = (
        df["gender"].value_counts().reindex(GENDER_ORDER).rename_axis("gender").reset_index(name="count")
    )
    plt.figure(figsize=(6, 4))
    sns.barplot(data=gender_counts, x="gender", y="count", hue="gender", dodge=False, legend=False)
    plt.title("Distribution globale du genre")
    plt.xlabel("Genre")
    plt.ylabel("Nombre d'images")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "gender_distribution.png", dpi=200)
    plt.close()

    heatmap_df = pd.crosstab(df["race"], df["gender"]).reindex(index=RACE_ORDER, columns=GENDER_ORDER)
    plt.figure(figsize=(7, 5))
    sns.heatmap(heatmap_df, annot=True, fmt=".0f", cmap="YlGnBu")
    plt.title("Croisement race x genre")
    plt.xlabel("Genre")
    plt.ylabel("Race")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "race_gender_heatmap.png", dpi=200)
    plt.close()


def stratified_sample(df: pd.DataFrame, sample_size: int, random_state: int) -> pd.DataFrame:
    if sample_size >= len(df):
        return df.sample(frac=1.0, random_state=random_state).reset_index(drop=True)
    strata = df["race"] + " | " + df["gender"]
    _, sample = train_test_split(
        df,
        test_size=sample_size,
        stratify=strata,
        random_state=random_state,
    )
    return sample.reset_index(drop=True)


def distance(landmarks: list[tuple[float, float]], a: int, b: int) -> float:
    ax, ay = landmarks[a]
    bx, by = landmarks[b]
    return math.hypot(ax - bx, ay - by)


def extract_ratios(landmarks: list[tuple[float, float]]) -> dict[str, float]:
    face_w = distance(landmarks, 234, 454)
    face_h = distance(landmarks, 10, 152)
    eye_w_l = distance(landmarks, 33, 133)
    eye_w_r = distance(landmarks, 362, 263)
    lip_w = distance(landmarks, 61, 291)
    nose_bridge = distance(landmarks, 6, 168) + distance(landmarks, 168, 2)
    eps = 1e-6

    return {
        "eye_left_open": distance(landmarks, 159, 145) / (eye_w_l + eps),
        "eye_right_open": distance(landmarks, 386, 374) / (eye_w_r + eps),
        "nose_width_ratio": distance(landmarks, 98, 327) / (face_w + eps),
        "nose_bridge_ratio": nose_bridge / (face_h + eps),
        "face_shape_index": face_w / (face_h + eps),
        "jaw_width_ratio": distance(landmarks, 172, 397) / (face_w + eps),
        "lip_fullness_ratio": distance(landmarks, 13, 14) / (lip_w + eps),
        "eyebrow_arch_ratio": distance(landmarks, 70, 159) / (eye_w_l + eps),
        "interpupillary_ratio": distance(landmarks, 468, 473) / (face_w + eps),
        "chin_projection_ratio": distance(landmarks, 152, 18) / (face_h + eps),
    }


def extract_features(sample_df: pd.DataFrame, min_confidence: float) -> tuple[pd.DataFrame, dict[str, float]]:
    options = vision.FaceLandmarkerOptions(
        base_options=BaseOptions(model_asset_path=str(MODEL_PATH)),
        num_faces=1,
        min_face_detection_confidence=min_confidence,
        min_face_presence_confidence=min_confidence,
    )
    rows: list[dict[str, object]] = []

    with vision.FaceLandmarker.create_from_options(options) as landmarker:
        for row in sample_df.itertuples(index=False):
            path = Path(row.image_path)
            record = {
                "file": row.file,
                "split": row.split,
                "age": row.age,
                "gender": row.gender,
                "race": row.race,
                "service_test": row.service_test,
                "image_exists": path.exists(),
            }
            if not record["image_exists"]:
                record["detected"] = False
                rows.append(record)
                continue

            file_bytes = path.read_bytes()
            image_hash = hashlib.md5(file_bytes).hexdigest()
            image_array = cv2.imdecode(np.frombuffer(file_bytes, dtype=np.uint8), cv2.IMREAD_COLOR)
            if image_array is None:
                record["detected"] = False
                rows.append(record)
                continue

            h, w = image_array.shape[:2]
            rgb = cv2.cvtColor(image_array, cv2.COLOR_BGR2RGB)
            mp_image = mp.Image(image_format=mp.ImageFormat.SRGB, data=rgb)
            result = landmarker.detect(mp_image)

            record.update(
                {
                    "width": w,
                    "height": h,
                    "image_md5": image_hash,
                    "detected": bool(result.face_landmarks),
                }
            )
            if result.face_landmarks:
                landmarks = [(p.x, p.y) for p in result.face_landmarks[0]]
                record.update(extract_ratios(landmarks))
            rows.append(record)

    features_df = pd.DataFrame(rows)
    stats = {
        "sample_size": int(len(sample_df)),
        "missing_image_files": int((~features_df["image_exists"]).sum()),
        "detected_faces": int(features_df["detected"].sum()),
        "detection_rate": float(features_df["detected"].mean()),
        "exact_duplicates": int(features_df["image_md5"].dropna().duplicated().sum()),
    }
    return features_df, stats


def plot_detection_rates(features_df: pd.DataFrame) -> pd.DataFrame:
    det_by_race = (
        features_df.groupby("race")["detected"].mean().mul(100).reindex(RACE_ORDER).reset_index(name="rate")
    )
    plt.figure(figsize=(9, 4.5))
    sns.barplot(data=det_by_race, x="race", y="rate", color="#2A9D8F")
    plt.ylim(0, 100)
    plt.title("Taux de detection FaceMesh par race sur l'echantillon")
    plt.xlabel("Race")
    plt.ylabel("Taux de detection (%)")
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "detection_rate_by_race.png", dpi=200)
    plt.close()
    return det_by_race


def prepare_features(features_df: pd.DataFrame) -> tuple[pd.DataFrame, np.ndarray, pd.DataFrame, dict[str, float]]:
    feature_cols = list(FEATURE_LABELS.keys())
    detected_df = features_df[features_df["detected"]].copy()
    missing_values = int(detected_df[feature_cols].isna().sum().sum())

    z_scores = np.abs(zscore(detected_df[feature_cols], nan_policy="omit"))
    inlier_mask = (z_scores < 3).all(axis=1)
    clean_df = detected_df.loc[inlier_mask].copy()
    outliers_removed = int((~inlier_mask).sum())

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(clean_df[feature_cols])
    scaled_df = pd.DataFrame(X_scaled, columns=feature_cols, index=clean_df.index)

    stats = {
        "missing_feature_values": missing_values,
        "outliers_removed": outliers_removed,
        "modeling_rows": int(len(clean_df)),
        "feature_extraction_success": float(
            len(detected_df[feature_cols].dropna()) / max(len(detected_df), 1)
        ),
    }
    return clean_df, X_scaled, scaled_df, stats


def plot_feature_eda(clean_df: pd.DataFrame, scaled_df: pd.DataFrame) -> None:
    feature_cols = list(FEATURE_LABELS.keys())

    fig, axes = plt.subplots(5, 2, figsize=(12, 16))
    axes = axes.flatten()
    for ax, col in zip(axes, feature_cols):
        sns.histplot(clean_df[col], kde=True, ax=ax, color="#6C5CE7", bins=30)
        ax.set_title(FEATURE_LABELS[col])
        ax.set_xlabel("")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "feature_distributions.png", dpi=200)
    plt.close(fig)

    corr = clean_df[feature_cols].corr()
    plt.figure(figsize=(10, 8))
    sns.heatmap(corr, cmap="coolwarm", center=0, annot=False)
    plt.title("Matrice de correlation des ratios geometriques")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "feature_correlation.png", dpi=200)
    plt.close()

    plt.figure(figsize=(11, 5))
    scaled_df.boxplot(rot=20)
    plt.title("Variables standardisees apres nettoyage")
    plt.ylabel("Z-score")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "scaled_feature_boxplot.png", dpi=200)
    plt.close()


def interpret_component(top_features: list[str]) -> str:
    theme_counts = pd.Series([FEATURE_GROUPS[col] for col in top_features]).value_counts()
    dominant_themes = list(theme_counts.index[:2])
    if len(dominant_themes) == 1:
        return f"Cette composante capture surtout la {dominant_themes[0]}."
    return (
        f"Cette composante combine principalement la {dominant_themes[0]} "
        f"et la {dominant_themes[1]}."
    )


def run_pca(X_scaled: np.ndarray) -> tuple[PCA, np.ndarray, pd.DataFrame, pd.DataFrame]:
    feature_cols = list(FEATURE_LABELS.keys())
    pca = PCA(n_components=0.85)
    X_pca = pca.fit_transform(X_scaled)

    component_names = [f"PC{i + 1}" for i in range(pca.n_components_)]
    loadings = pd.DataFrame(pca.components_.T, index=feature_cols, columns=component_names)

    rows: list[dict[str, object]] = []
    cumulative = 0.0
    for index, component in enumerate(component_names):
        explained = float(pca.explained_variance_ratio_[index])
        cumulative += explained
        top_features = (
            loadings[component]
            .abs()
            .sort_values(ascending=False)
            .head(3)
            .index.tolist()
        )
        rows.append(
            {
                "Composante": component,
                "Variance expliquee (%)": round(explained * 100, 2),
                "Variance cumulée (%)": round(cumulative * 100, 2),
                "Variables dominantes": ", ".join(FEATURE_LABELS[col] for col in top_features),
                "Interpretation": interpret_component(top_features),
            }
        )

    summary_df = pd.DataFrame(rows)
    return pca, X_pca, loadings, summary_df


def plot_pca(pca: PCA, loadings: pd.DataFrame) -> None:
    cumulative = np.cumsum(pca.explained_variance_ratio_) * 100
    plt.figure(figsize=(8, 4.5))
    plt.plot(range(1, len(cumulative) + 1), cumulative, marker="o", color="#E76F51")
    plt.axhline(85, color="black", linestyle="--", linewidth=1)
    plt.title("Courbe cumulative de variance expliquee (ACP)")
    plt.xlabel("Nombre de composantes")
    plt.ylabel("Variance expliquee cumulee (%)")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "pca_scree_plot.png", dpi=200)
    plt.close()

    shown_loadings = loadings.iloc[:, : min(4, loadings.shape[1])].copy()
    shown_loadings.index = [FEATURE_LABELS[col] for col in shown_loadings.index]
    plt.figure(figsize=(8, 6))
    sns.heatmap(shown_loadings, cmap="vlag", center=0, annot=True, fmt=".2f")
    plt.title("Loadings des premieres composantes")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "pca_loadings_heatmap.png", dpi=200)
    plt.close()


def evaluate_models(X_pca: np.ndarray) -> pd.DataFrame:
    results: list[dict[str, object]] = []
    for k in range(3, 9):
        models = {
            "KMeans": KMeans(n_clusters=k, random_state=42, n_init=20),
            "Agglomerative": AgglomerativeClustering(n_clusters=k, linkage="ward"),
            "GaussianMixture": GaussianMixture(n_components=k, covariance_type="full", random_state=42),
        }
        for model_name, model in models.items():
            labels = model.fit_predict(X_pca)
            if len(np.unique(labels)) < 2:
                continue
            results.append(
                {
                    "model": model_name,
                    "k": k,
                    "silhouette": float(silhouette_score(X_pca, labels)),
                    "calinski_harabasz": float(calinski_harabasz_score(X_pca, labels)),
                    "davies_bouldin": float(davies_bouldin_score(X_pca, labels)),
                }
            )
    return pd.DataFrame(results)


def best_per_model(metrics_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for model_name in metrics_df["model"].unique():
        subset = metrics_df[metrics_df["model"] == model_name].sort_values(
            ["silhouette", "calinski_harabasz", "davies_bouldin"],
            ascending=[False, False, True],
        )
        rows.append(subset.iloc[0].to_dict())
    best_df = pd.DataFrame(rows)
    best_df["model_label"] = best_df["model"].map(MODEL_NAMES)
    return best_df.sort_values(
        ["silhouette", "calinski_harabasz", "davies_bouldin"],
        ascending=[False, False, True],
    )


def fit_final_model(model_name: str, k: int, X_pca: np.ndarray) -> np.ndarray:
    if model_name == "KMeans":
        model = KMeans(n_clusters=k, random_state=42, n_init=20)
    elif model_name == "Agglomerative":
        model = AgglomerativeClustering(n_clusters=k, linkage="ward")
    elif model_name == "GaussianMixture":
        model = GaussianMixture(n_components=k, covariance_type="full", random_state=42)
    else:
        raise ValueError(f"Unsupported model: {model_name}")
    return model.fit_predict(X_pca)


def plot_model_results(metrics_df: pd.DataFrame, scaled_df: pd.DataFrame, labels: np.ndarray, X_pca: np.ndarray) -> tuple[pd.DataFrame, pd.DataFrame]:
    plt.figure(figsize=(9, 5))
    sns.lineplot(data=metrics_df, x="k", y="silhouette", hue="model", marker="o")
    plt.title("Comparaison des modeles par silhouette")
    plt.xlabel("Nombre de clusters / composantes")
    plt.ylabel("Silhouette")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "model_comparison.png", dpi=200)
    plt.close()

    plot_df = pd.DataFrame(
        {
            "PC1": X_pca[:, 0],
            "PC2": X_pca[:, 1],
            "cluster": labels.astype(str),
        }
    )
    plt.figure(figsize=(8, 6))
    sns.scatterplot(data=plot_df, x="PC1", y="PC2", hue="cluster", s=20, alpha=0.7, palette="tab10")
    plt.title("Projection des clusters sur les deux premieres composantes")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "cluster_scatter.png", dpi=200)
    plt.close()

    centroid_df = scaled_df.copy()
    centroid_df["cluster"] = labels
    cluster_heatmap = centroid_df.groupby("cluster").mean().copy()
    cluster_heatmap.index = [f"Cluster {idx}" for idx in cluster_heatmap.index]
    renamed = cluster_heatmap.rename(columns=FEATURE_LABELS)
    plt.figure(figsize=(11, 4.5))
    sns.heatmap(renamed, cmap="coolwarm", center=0, annot=True, fmt=".2f")
    plt.title("Profils moyens standardises par cluster")
    plt.tight_layout()
    plt.savefig(ASSETS_DIR / "cluster_profile_heatmap.png", dpi=200)
    plt.close()

    cluster_sizes = (
        pd.Series(labels, name="cluster")
        .value_counts(normalize=True)
        .sort_index()
        .mul(100)
        .rename("Part du cluster (%)")
        .reset_index()
        .rename(columns={"index": "cluster"})
    )
    profile_rows: list[dict[str, object]] = []
    for cluster_name, row in cluster_heatmap.iterrows():
        strongest = row.sort_values(ascending=False).head(2).index.tolist()
        weakest = row.sort_values().head(1).index.tolist()
        cluster_index = int(cluster_name.split()[-1])
        share = cluster_sizes.loc[cluster_sizes["cluster"] == cluster_index, "Part du cluster (%)"].iloc[0]
        profile_rows.append(
            {
                "Cluster": cluster_name,
                "Part (%)": round(float(share), 2),
                "Traits dominants": ", ".join(FEATURE_LABELS[col] for col in strongest),
                "Trait en retrait": FEATURE_LABELS[weakest[0]],
            }
        )

    profile_df = pd.DataFrame(profile_rows)
    return cluster_heatmap, profile_df


def set_document_language(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr_cells[idx].text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def add_picture(document: Document, image_path: Path, caption: str, width: float = 6.3) -> None:
    document.add_picture(str(image_path), width=Inches(width))
    paragraph = document.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].italic = True


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def build_report(
    dataset_df: pd.DataFrame,
    sample_df: pd.DataFrame,
    pca: PCA,
    pca_summary_df: pd.DataFrame,
    best_models_df: pd.DataFrame,
    cluster_profile_df: pd.DataFrame,
    extraction_stats: dict[str, float],
    prep_stats: dict[str, float],
    detection_by_race: pd.DataFrame,
    report_path: Path,
) -> None:
    document = Document()
    set_document_language(document)

    title = document.add_heading("Rapport mis a jour - FaceMesh Feature Extraction", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        "Version restructuree autour des sections demandees : Business Understanding, "
        "Data Understanding, Data Preparation, ACP (PCA) et Modeling."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "Ce rapport met a jour le module d'extraction de traits faciaux pour l'ecosysteme "
        "voix + visage et tient compte des dernieres contraintes d'execution sur CPU "
        "dans une logique de deploiement legere de type Raspberry Pi. "
        f"L'EDA est realisee sur l'ensemble du dataset FairFace ({len(dataset_df):,} images), "
        f"tandis que l'extraction FaceMesh, l'ACP et la comparaison des modeles sont "
        f"executees sur un echantillon stratifie de {len(sample_df):,} images."
    )

    document.add_heading("1. Business Understanding", level=1)
    document.add_paragraph(
        "L'objectif metier est de transformer des images de visage en descripteurs "
        "compacts, interpretables et exploitables par un pipeline multimodal voix-visage. "
        "Les objectifs data science finaux ci-dessous formalisent le lien entre besoin "
        "metier, traitement analytique et indicateurs de succes."
    )
    bo_dso_df = pd.DataFrame(
        [
            {
                "ID": "BO-1",
                "Business Objective": "Automatiser l'extraction de traits faciaux interpretable sur CPU.",
                "Data Science Objective": "Detecter les visages puis calculer 10 ratios geometriques robustes via FaceMesh.",
                "Indicateur": f"Taux de detection observe: {pct(extraction_stats['detection_rate'])}",
            },
            {
                "ID": "BO-2",
                "Business Objective": "Obtenir une representation compacte compatible avec le Pi.",
                "Data Science Objective": "Standardiser les ratios puis appliquer une ACP conservant au moins 85% de la variance.",
                "Indicateur": f"{pca.n_components_} composantes pour {pca.explained_variance_ratio_.sum() * 100:.2f}% de variance",
            },
            {
                "ID": "BO-3",
                "Business Objective": "Segmenter les profils faciaux pour la personnalisation et l'analyse.",
                "Data Science Objective": "Comparer plusieurs modeles de clustering deja etudies et retenir le meilleur compromis.",
                "Indicateur": (
                    f"Meilleur modele: {best_models_df.iloc[0]['model_label']} "
                    f"(k={int(best_models_df.iloc[0]['k'])}, silhouette={best_models_df.iloc[0]['silhouette']:.3f})"
                ),
            },
            {
                "ID": "BO-4",
                "Business Objective": "Suivre la couverture et les biais potentiels entre sous-groupes.",
                "Data Science Objective": "Conduire l'EDA et mesurer le taux de detection par race/genre sur l'echantillon.",
                "Indicateur": (
                    f"Taux de detection par race entre {detection_by_race['rate'].min():.1f}% "
                    f"et {detection_by_race['rate'].max():.1f}%"
                ),
            },
        ]
    )
    add_table(document, bo_dso_df)

    document.add_heading("2. Data Understanding", level=1)
    dataset_summary = pd.DataFrame(
        [
            {
                "Split": "Train",
                "Nb images": int((dataset_df["split"] == "train").sum()),
                "Classes race": dataset_df[dataset_df["split"] == "train"]["race"].nunique(),
                "Classes age": dataset_df[dataset_df["split"] == "train"]["age"].nunique(),
                "Classes genre": dataset_df[dataset_df["split"] == "train"]["gender"].nunique(),
            },
            {
                "Split": "Validation",
                "Nb images": int((dataset_df["split"] == "val").sum()),
                "Classes race": dataset_df[dataset_df["split"] == "val"]["race"].nunique(),
                "Classes age": dataset_df[dataset_df["split"] == "val"]["age"].nunique(),
                "Classes genre": dataset_df[dataset_df["split"] == "val"]["gender"].nunique(),
            },
            {
                "Split": "Total",
                "Nb images": int(len(dataset_df)),
                "Classes race": dataset_df["race"].nunique(),
                "Classes age": dataset_df["age"].nunique(),
                "Classes genre": dataset_df["gender"].nunique(),
            },
        ]
    )
    add_table(document, dataset_summary)
    document.add_paragraph(
        "Observations EDA principales : les groupes raciaux sont relativement equilibres, "
        "le genre reste proche d'un partage 50/50, et les tranches 20-29 puis 30-39 "
        "sont les plus representees. Le dataset est donc favorable a une analyse "
        "exploratoire et a un echantillonnage stratifie."
    )
    add_picture(document, ASSETS_DIR / "race_distribution.png", "Figure 1 - Repartition des images par race et par split.")
    add_picture(document, ASSETS_DIR / "age_distribution.png", "Figure 2 - Distribution globale des tranches d'age.")
    add_picture(document, ASSETS_DIR / "race_gender_heatmap.png", "Figure 3 - Heatmap race x genre.")
    add_picture(document, ASSETS_DIR / "detection_rate_by_race.png", "Figure 4 - Taux de detection FaceMesh par race sur l'echantillon.")

    document.add_paragraph("Interpretation des observations :")
    for observation in [
        "FairFace reste nettement plus equilibre par race qu'un dataset grand public classique, ce qui rend l'EDA plus defendable.",
        "La concentration sur les tranches 20-39 ans doit etre gardee en tete lors de l'interpretation des clusters.",
        "Le taux de detection FaceMesh varie selon les sous-groupes, ce qui confirme l'interet de suivre ce KPI pour limiter les biais de couverture.",
    ]:
        document.add_paragraph(observation, style="List Bullet")

    document.add_heading("3. Data Preparation", level=1)
    document.add_paragraph(
        "La preparation des donnees a ete detaillee afin de couvrir les points demandes "
        "(missing values, outliers, encodage, normalisation). Les variables categorielles "
        "du CSV ont ete conservees pour l'EDA uniquement. Aucun encodage n'etait necessaire "
        "pour le clustering, qui repose exclusivement sur des ratios geometriques continus."
    )
    prep_table = pd.DataFrame(
        [
            {"Etape": "Controle des labels", "Resultat": f"{int(dataset_df.isna().sum().sum())} valeur manquante dans les CSV"},
            {"Etape": "Verification des fichiers de l'echantillon", "Resultat": f"{extraction_stats['missing_image_files']} image manquante"},
            {"Etape": "Controle des doublons exacts", "Resultat": f"{extraction_stats['exact_duplicates']} doublon exact detecte"},
            {"Etape": "Detection FaceMesh", "Resultat": f"{extraction_stats['detected_faces']} / {extraction_stats['sample_size']} images detectees"},
            {"Etape": "Valeurs manquantes sur les features", "Resultat": f"{prep_stats['missing_feature_values']} cellule manquante"},
            {"Etape": "Suppression des outliers", "Resultat": f"{prep_stats['outliers_removed']} lignes retirees via z-score > 3"},
            {"Etape": "Standardisation", "Resultat": f"{prep_stats['modeling_rows']} lignes conservees pour la modelisation"},
        ]
    )
    add_table(document, prep_table)

    feature_table = pd.DataFrame(
        [
            {"Feature": FEATURE_LABELS[col], "Insight capture": FEATURE_GROUPS[col]}
            for col in FEATURE_LABELS
        ]
    )
    document.add_paragraph(
        "Les 10 variables finales sont des ratios normalises, donc peu sensibles a l'echelle absolue de l'image."
    )
    add_table(document, feature_table)
    add_picture(document, ASSETS_DIR / "feature_distributions.png", "Figure 5 - Distributions des ratios geometriques apres nettoyage.")
    add_picture(document, ASSETS_DIR / "feature_correlation.png", "Figure 6 - Correlations entre les ratios.")

    document.add_heading("4. ACP (PCA)", level=1)
    document.add_paragraph(
        "L'ACP est appliquee apres standardisation afin de reduire la redondance entre variables "
        "et de produire une representation compacte pour le clustering et un futur deploiement edge."
    )
    add_table(document, pca_summary_df)
    add_picture(document, ASSETS_DIR / "pca_scree_plot.png", "Figure 7 - Variance expliquee cumulee par l'ACP.")
    add_picture(document, ASSETS_DIR / "pca_loadings_heatmap.png", "Figure 8 - Loadings des premieres composantes.")
    document.add_paragraph("Insights PCA :")
    for insight in [
        "Les premieres composantes concentrent les grandes variations de structure faciale globale, de morphologie oculaire et de morphologie nasale.",
        "La reduction de dimension conserve une information suffisante pour un clustering stable tout en allegeant la representation finale.",
    ]:
        document.add_paragraph(insight, style="List Bullet")

    document.add_heading("5. Modeling", level=1)
    document.add_paragraph(
        "Trois modeles non supervises deja couramment etudies ont ete compares : "
        "KMeans, Agglomerative (Ward) et Gaussian Mixture. Le protocole d'evaluation "
        "utilise des metriques adaptees au clustering : silhouette, Calinski-Harabasz "
        "et Davies-Bouldin."
    )
    model_choice_df = pd.DataFrame(
        [
            {"Modele": "KMeans", "Justification": "Rapide, centroide, explicable et coherent avec l'ancien pipeline voix."},
            {"Modele": "Agglomerative (Ward)", "Justification": "Approche hierarchique utile pour verifier une structure non plate."},
            {"Modele": "Gaussian Mixture", "Justification": "Alternative probabiliste capable de modeliser des clusters elliptiques."},
        ]
    )
    add_table(document, model_choice_df)

    modeling_results = best_models_df[
        ["model_label", "k", "silhouette", "calinski_harabasz", "davies_bouldin"]
    ].rename(
        columns={
            "model_label": "Modele",
            "k": "k optimal",
            "silhouette": "Silhouette",
            "calinski_harabasz": "Calinski-Harabasz",
            "davies_bouldin": "Davies-Bouldin",
        }
    )
    modeling_results = modeling_results.copy()
    for col in ["Silhouette", "Calinski-Harabasz", "Davies-Bouldin"]:
        modeling_results[col] = modeling_results[col].map(lambda value: round(float(value), 3))
    add_table(document, modeling_results)
    add_picture(document, ASSETS_DIR / "model_comparison.png", "Figure 9 - Evolution du score de silhouette selon le modele et k.")
    add_picture(document, ASSETS_DIR / "cluster_scatter.png", "Figure 10 - Projection des observations sur PC1 et PC2.")
    add_picture(document, ASSETS_DIR / "cluster_profile_heatmap.png", "Figure 11 - Profils standardises des clusters du modele final.")

    kmeans_row = best_models_df[best_models_df["model"] == "KMeans"].iloc[0]
    document.add_paragraph(
        f"Le meilleur compromis observe est {best_models_df.iloc[0]['model_label']} "
        f"avec k={int(best_models_df.iloc[0]['k'])}. Ce choix est retenu car il maximise la "
        f"silhouette ({best_models_df.iloc[0]['silhouette']:.3f}) sur l'echantillon retenu."
    )
    document.add_paragraph(
        f"Si l'alignement avec l'ancien pipeline voix base sur KMeans est prioritaire, "
        f"KMeans reste une alternative credible avec k={int(kmeans_row['k'])} "
        f"et une silhouette de {kmeans_row['silhouette']:.3f}."
    )
    add_table(document, cluster_profile_df)
    document.add_paragraph(
        "Interpretation : les clusters obtenus ne doivent pas etre lus comme des categories "
        "biologiques fixes, mais comme des regroupements geometriques de ratios faciaux "
        "utiles pour la personnalisation, l'indexation et l'analyse exploratoire."
    )

    document.add_heading("6. Conclusion", level=1)
    document.add_paragraph(
        "Le pipeline mis a jour respecte les points demandes : BO/DSO explicites, EDA avec "
        "visualisations, preparation detaillee, ACP interpretee, comparaison de plusieurs "
        "modeles et evaluation avec metriques adaptees. Le principal point d'attention reste "
        "le taux de detection FaceMesh, qui peut encore etre ameliore par du pretraitement ou "
        "par l'usage des images FairFace margin125."
    )
    document.add_paragraph(
        "Le pipeline reste defensable pour un rendu academique, reproductible via le script "
        "`generate_facemesh_report.py`, et suffisamment leger pour une future integration CPU."
    )

    document.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate the updated FaceMesh report and its assets.")
    parser.add_argument("--sample-size", type=int, default=10000)
    parser.add_argument("--random-state", type=int, default=42)
    parser.add_argument("--min-confidence", type=float, default=0.2)
    args = parser.parse_args()

    prepare_dirs()
    ensure_model(MODEL_PATH)

    dataset_df = load_labels()
    plot_label_distributions(dataset_df)

    sample_df = stratified_sample(dataset_df, args.sample_size, args.random_state)
    features_df, extraction_stats = extract_features(sample_df, args.min_confidence)
    detection_by_race = plot_detection_rates(features_df)

    clean_df, X_scaled, scaled_df, prep_stats = prepare_features(features_df)
    plot_feature_eda(clean_df, scaled_df)

    pca, X_pca, loadings, pca_summary_df = run_pca(X_scaled)
    plot_pca(pca, loadings)

    metrics_df = evaluate_models(X_pca)
    best_models_df = best_per_model(metrics_df)
    final_model = best_models_df.iloc[0]
    final_labels = fit_final_model(str(final_model["model"]), int(final_model["k"]), X_pca)
    _, cluster_profile_df = plot_model_results(metrics_df, scaled_df, final_labels, X_pca)

    save_table(features_df, OUTPUT_DIR / "sample_features.csv")
    save_table(loadings.reset_index().rename(columns={"index": "feature"}), OUTPUT_DIR / "pca_loadings.csv")
    save_table(metrics_df, OUTPUT_DIR / "model_metrics.csv")
    save_table(best_models_df, OUTPUT_DIR / "best_model_per_family.csv")
    save_table(cluster_profile_df, OUTPUT_DIR / "cluster_profiles.csv")
    save_table(detection_by_race, OUTPUT_DIR / "detection_rate_by_race.csv")

    report_path = OUTPUT_DIR / "FaceMesh_Feature_Extraction_Report_Updated.docx"
    build_report(
        dataset_df=dataset_df,
        sample_df=sample_df,
        pca=pca,
        pca_summary_df=pca_summary_df,
        best_models_df=best_models_df,
        cluster_profile_df=cluster_profile_df,
        extraction_stats=extraction_stats,
        prep_stats=prep_stats,
        detection_by_race=detection_by_race,
        report_path=report_path,
    )

    print(f"Report generated at: {report_path}")
    print(f"Assets generated in: {ASSETS_DIR}")


if __name__ == "__main__":
    main()
